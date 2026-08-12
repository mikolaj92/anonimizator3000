import base64
import html
import json
import re
import time
from datetime import date
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from doctotext import DOCX_MIME
from docx import Document
from fastapi.testclient import TestClient
from itsdangerous import TimestampSigner
from my_usermanager.models import Permission
from my_usermanager.permissions import ADMIN_ROLE_NAME
from my_usermanager.sessions import SessionPrincipal, write_session_principal
from posejdon import TextAnonymizer

import anonimizator3000.main as main_module
from anonimizator3000.main import _attachment_header, _settings_boot, app
from anonimizator3000.upload import UploadError, read_multipart_document


@pytest.fixture(autouse=True)
def _use_regex_only_anonymizer(monkeypatch, request) -> None:
    """UI tests do not require detector startup except for its regression test."""
    if request.node.get_closest_marker("real_anonymizer") is None:
        monkeypatch.setattr(main_module, "create_anonymizer", lambda settings: TextAnonymizer())


def _authenticated_session() -> dict:
    session = {"user": {"id": "audit-user", "name": "audit_user", "is_admin": True}}
    write_session_principal(
        session,
        SessionPrincipal(
            user_id="audit-user",
            username="audit_user",
            display_name="Audit User",
            roles=frozenset({ADMIN_ROLE_NAME}),
            permissions=frozenset({Permission("admin.access")}),
        ),
    )
    return session


@pytest.mark.real_anonymizer
def test_default_lifespan_anonymizes_gliner_only_pii_with_real_stack() -> None:
    with TestClient(app) as client:
        assert client.get("/healthz").status_code == 200
        result = client.app.state.queue._anonymizer.anonymize_segments(
            ["Spotkanie z Mikołajem Brzęczyszczykiewiczem."], replacement_style="mask"
        )

    assert "Mikołajem Brzęczyszczykiewiczem" not in result.texts[0]
    assert result.findings["PERSON"] >= 1


def _docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _docx_bytes_with_hyperlink(text: str) -> bytes:
    """A DOCX whose body references the ``r:`` namespace, like real Word files."""
    with ZipFile(BytesIO(_docx_bytes(text))) as archive:
        parts = {name: archive.read(name) for name in archive.namelist()}
    document_xml = parts["word/document.xml"].decode()
    parts["word/document.xml"] = document_xml.replace(
        f"<w:r><w:t>{text}</w:t></w:r>",
        f'<w:hyperlink r:id="rId999"><w:r><w:t>{text}</w:t></w:r></w:hyperlink>',
    ).encode()
    output = BytesIO()
    with ZipFile(output, "w") as archive:
        for name, data in parts.items():
            archive.writestr(name, data)
    return output.getvalue()


def test_index_does_not_show_removed_header_copy() -> None:
    with TestClient(app) as client:
        response = client.get("/")

    assert response.status_code == 200
    assert 'aria-label="Informacje o prywatności"' in response.text
    assert "Dokument nie jest zapisywany na dysku" in response.text
    assert "wynik też wygasa" in response.text
    assert "Lokalnie. Bez zapisu." not in response.text
    assert "Anonimizator3000" not in response.text
    assert "in-memory" not in response.text
    assert "Gotowy dokument pojawi się tutaj." not in response.text
    assert "cdn.jsdelivr.net" not in response.text
    assert "unpkg.com" not in response.text
    # Platform assets from app-factory (same-origin /static/platform), not vendored.
    assert "/static/platform/" in response.text
    assert "/static/basecoat/" not in response.text
    assert "/static/htmx.min.js" not in response.text
    assert "/static/app.css" in response.text
    assert "100 dokumentów" in response.text
    assert "10 minut" in response.text


def test_docx_upload_poll_and_download_flow_returns_docx() -> None:
    with TestClient(app) as client:
        response = client.post(
            "/jobs",
            headers={"HX-Request": "true"},
            files={
                "document": (
                    "sample.docx",
                    _docx_bytes("Jan Kowalski, PESEL 44051401359, email jan@example.com"),
                    DOCX_MIME,
                )
            },
        )

        assert response.status_code == 200
        assert 'role="progressbar"' in response.text
        assert 'id="job-card"' in response.text
        assert 'id="job-panel"' not in response.text
        job_id = response.text.split("/jobs/", 1)[1].split('"', 1)[0]

        for _ in range(50):
            status_response = client.get(
                f"/jobs/{job_id}", headers={"HX-Request": "true"}
            )
            assert status_response.status_code == 200
            if "Gotowe" in status_response.text:
                assert "textarea" not in status_response.text
                assert "Pobierz" in status_response.text
                break
            time.sleep(0.05)
        else:
            raise AssertionError("Job did not finish")

        download = client.get(f"/jobs/{job_id}/download")

        assert download.status_code == 200
        assert download.headers["content-type"] == DOCX_MIME
        downloaded = Document(BytesIO(download.content))
        text = "\n".join(paragraph.text for paragraph in downloaded.paragraphs)
        assert "Jan Kowalski" not in text
        assert "44051401359" not in text
        assert "jan@example.com" not in text


def test_docx_download_keeps_root_namespace_declarations() -> None:
    with TestClient(app) as client:
        response = client.post(
            "/jobs",
            headers={"HX-Request": "true"},
            files={
                "document": (
                    "sample.docx",
                    _docx_bytes_with_hyperlink("Jan Kowalski"),
                    DOCX_MIME,
                )
            },
        )
        job_id = response.text.split("/jobs/", 1)[1].split('"', 1)[0]

        for _ in range(50):
            status = client.get(f"/jobs/{job_id}", headers={"HX-Request": "true"})
            if "Gotowe" in status.text:
                break
            time.sleep(0.05)
        else:
            raise AssertionError("Job did not finish")

        download = client.get(f"/jobs/{job_id}/download")

    assert download.headers["content-type"] == DOCX_MIME
    document_xml = ZipFile(BytesIO(download.content)).read("word/document.xml").decode()
    start = document_xml.index("<w:document")
    root = document_xml[start : document_xml.index(">", start) + 1]
    # mc:Ignorable references these prefixes; dropping their xmlns breaks Word.
    assert "xmlns:w14=" in root
    assert "xmlns:wp14=" in root
    # The file must still open as a valid DOCX.
    assert "Jan Kowalski" not in "".join(
        p.text for p in Document(BytesIO(download.content)).paragraphs
    )


def test_healthz_returns_only_update_date() -> None:
    with TestClient(app) as client:
        response = client.get("/healthz")

    assert response.status_code == 200
    assert date.fromisoformat(response.text.strip())
    assert response.text.count("\n") == 1


def test_plain_form_upload_redirects_to_server_rendered_job_page() -> None:
    with TestClient(app) as client:
        response = client.post(
            "/jobs",
            files={"document": ("sample.txt", b"hello", "text/plain")},
            follow_redirects=False,
        )
        assert response.status_code == 303
        page = client.get(response.headers["location"])

    assert page.status_code == 200
    assert '<form\n        hx-post="/jobs"' in page.text
    assert 'id="job-card"' in page.text
    assert page.text.count('id="job-panel"') == 1


def test_plain_form_validation_error_returns_full_page() -> None:
    with TestClient(app) as client:
        response = client.post(
            "/jobs",
            files={"document": ("big.txt", b"x" * 5_200_000, "text/plain")},
        )

    assert response.status_code == 413
    assert "Odrzucono upload" in response.text
    assert 'id="job-panel"' in response.text
    assert '<form\n        hx-post="/jobs"' in response.text


def test_upload_size_limit_returns_fragment() -> None:
    with TestClient(app) as client:
        response = client.post(
            "/jobs",
            headers={"HX-Request": "true"},
            files={"document": ("big.txt", b"x" * 5_200_000, "text/plain")},
        )

        assert response.status_code == 413
        assert "Odrzucono upload" in response.text
        assert 'id="job-panel"' not in response.text


def test_htmx_config_swaps_client_error_fragments() -> None:
    with TestClient(app) as client:
        page = client.get("/")
        missing_job = client.get(
            "/jobs/not-a-job", headers={"HX-Request": "true"}
        )

    match = re.search(
        r'<meta\s+name="htmx-config"\s+content=\'([^\']+)\'', page.text
    )
    assert match is not None
    response_handling = json.loads(html.unescape(match.group(1)))["responseHandling"]

    for status_code in (413, missing_job.status_code):
        handling = next(
            rule
            for rule in response_handling
            if re.fullmatch(rule["code"], str(status_code))
        )
        assert handling == {"code": "4..", "swap": True, "error": True}

    assert missing_job.status_code == 404
    assert "Zadanie wygasło albo nie istnieje." in missing_job.text


@pytest.mark.asyncio
async def test_invalid_content_length_returns_upload_error() -> None:
    request = _FakeUploadRequest(content_length="not-a-number")

    with pytest.raises(UploadError, match="Content-Length") as error:
        await read_multipart_document(request, max_file_bytes=100, max_body_bytes=200)

    assert error.value.status_code == 400


def test_attachment_header_uses_encoded_filename() -> None:
    header = _attachment_header('..\\zażółć"\r\nx.docx')

    assert header == (
        "attachment; filename*=UTF-8''za%C5%BC%C3%B3%C5%82%C4%87%22%0D%0Ax.docx"
    )
    assert 'filename="' not in header


class _FakeUploadRequest:
    def __init__(self, *, content_length: str) -> None:
        self.headers = {
            "content-type": "multipart/form-data; boundary=x",
            "content-length": content_length,
        }

    async def stream(self):
        yield b""


def test_platform_auth_routes_exist() -> None:
    """Package-owned auth/account/admin surfaces are reachable anonymously."""
    with TestClient(app) as client:
        login = client.get("/login")
        register = client.get("/register")
        account = client.get("/account", follow_redirects=False)
        profile = client.post("/account/profile", follow_redirects=False)
        admin_users = client.get("/admin/users", follow_redirects=False)
        logout = client.post("/logout", follow_redirects=False)

    assert login.status_code == 200
    assert register.status_code == 200
    assert account.status_code == 303
    assert profile.status_code == 401
    assert admin_users.status_code == 303
    assert logout.status_code == 303


def test_legacy_untyped_session_does_not_authenticate() -> None:
    legacy_session = {
        "user": {"id": "audit-user", "name": "audit_user", "is_admin": True}
    }
    payload = base64.b64encode(json.dumps(legacy_session).encode())
    cookie = TimestampSigner(_settings_boot.session_secret).sign(payload).decode()

    with TestClient(app) as client:
        client.cookies.set(_settings_boot.session_cookie_name, cookie)
        account = client.get("/account", follow_redirects=False)
        admin_users = client.get("/admin/users", follow_redirects=False)

    assert account.status_code == 303
    assert admin_users.status_code == 303


@pytest.mark.parametrize("path", ("/login", "/account", "/admin/users"))
def test_platform_ui_shell_loads_same_origin_stack(path: str) -> None:
    session = _authenticated_session()
    payload = base64.b64encode(json.dumps(session).encode())
    cookie = TimestampSigner(_settings_boot.session_secret).sign(payload).decode()

    with TestClient(app) as client:
        client.cookies.set(_settings_boot.session_cookie_name, cookie)
        response = client.get(path)

    assert response.status_code == 200
    for asset in (
        "/static/platform/basecoat-factory.min.css",
        "/static/platform/basecoat-js.min.js",
        "/static/platform/htmx.min.js",
        "/static/platform/alpine.min.js",
    ):
        assert asset in response.text
    assert "cdn.jsdelivr.net" not in response.text
    assert "unpkg.com" not in response.text


def test_platform_asset_is_served_same_origin() -> None:
    with TestClient(app) as client:
        response = client.get("/static/platform/basecoat-factory.min.css")

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/css")


def test_host_static_contains_only_product_assets() -> None:
    static_dir = Path(__file__).resolve().parents[1] / "src/anonimizator3000/static"

    assert {
        path.relative_to(static_dir).as_posix()
        for path in static_dir.rglob("*")
        if path.is_file()
    } == {"app.css"}


def test_landing_disclosure_is_native_and_has_an_accessible_relationship() -> None:
    with TestClient(app) as client:
        response = client.get("/")

    assert response.status_code == 200
    assert '<details id="privacy-disclosure" class="info-disclosure">' in response.text
    assert 'aria-controls="privacy-information"' in response.text
    assert '<div id="privacy-information" class="info-panel">' in response.text

    templates_dir = (
        Path(__file__).resolve().parents[1] / "src/anonimizator3000/templates"
    )
    host_templates = "\n".join(
        path.read_text(encoding="utf-8") for path in templates_dir.rglob("*.html")
    )
    assert "$store" not in host_templates
    assert "Alpine.store" not in host_templates


def test_base_delegates_platform_chrome_to_product_shell() -> None:
    template = Path(__file__).resolve().parents[1] / "src/anonimizator3000/templates/base.html"
    base = template.read_text(encoding="utf-8")

    assert 'extends "app_factory/product_shell.html"' in base
    assert "basecoat/basecoat" not in base
    assert "htmx.min.js" not in base
    assert "head_assets" not in base
    assert "platform_theme_locale" not in base
    assert "platform_sidebar" not in base
    assert "platform_session" not in base


@pytest.mark.parametrize("path", ("/", "/account", "/admin/users"))
def test_product_surfaces_share_platform_chrome_contract(path: str) -> None:
    session = _authenticated_session()
    payload = base64.b64encode(json.dumps(session).encode())
    cookie = TimestampSigner(_settings_boot.session_secret).sign(payload).decode()

    with TestClient(app) as client:
        client.cookies.set(_settings_boot.session_cookie_name, cookie)
        response = client.get(f"{path}?lang=de")

    assert response.status_code == 200
    assert '<html lang="de">' in response.text
    for marker in (
        'id="sidebar"',
        'id="sidebar-toggle"',
        "data-platform-theme-locale",
        'data-href="/?lang=pl"' if path == "/" else f'data-href="{path}?lang=pl"',
        "/static/platform/basecoat-factory.min.css",
        "/static/platform/basecoat-js.min.js",
        "/static/platform/htmx.min.js",
        "/static/platform/alpine.min.js",
    ):
        assert marker in response.text


def test_invalid_landing_locale_falls_back_to_default() -> None:
    with TestClient(app) as client:
        response = client.get("/?lang=not-supported")

    assert response.status_code == 200
    assert '<html lang="pl">' in response.text
    assert '<option\n      value="pl"\n      data-href="/?lang=pl"\n      selected' in response.text


def test_product_shell_uses_platform_controls_without_logout() -> None:
    session = _authenticated_session()
    payload = base64.b64encode(json.dumps(session).encode())
    cookie = TimestampSigner(_settings_boot.session_secret).sign(payload).decode()

    with TestClient(app) as client:
        client.cookies.set(_settings_boot.session_cookie_name, cookie)
        response = client.get("/")

    assert response.status_code == 200
    assert "data-platform-theme-locale" in response.text
    assert "data-platform-account-link" in response.text
    assert 'action="/logout"' not in response.text


def test_logout_is_only_rendered_on_account_surface() -> None:
    session = _authenticated_session()
    payload = base64.b64encode(json.dumps(session).encode())
    cookie = TimestampSigner(_settings_boot.session_secret).sign(payload).decode()

    with TestClient(app) as client:
        client.cookies.set(_settings_boot.session_cookie_name, cookie)
        response = client.get("/account")

    assert response.status_code == 200
    assert "data-platform-session" in response.text
    assert 'action="/logout"' in response.text
