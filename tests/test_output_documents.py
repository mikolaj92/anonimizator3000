import re
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import fitz
import pytest
from doctotext import DOCX_MIME, PDF_MIME, DocumentError, load_document
from docx import Document
from posejdon import TextAnonymizer

from anonimizator3000.pipeline import process_document

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

UNICODE_FONT_CANDIDATES = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    "/usr/local/share/fonts/DejaVuSans.ttf",
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
)


def _process(filename: str, content_type: str, data: bytes, *, max_text_chars: int = 10_000):
    return process_document(
        filename,
        content_type,
        data,
        anonymizer=TextAnonymizer(),
        max_text_chars=max_text_chars,
    )


def _docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _unicode_pdf_bytes(*pages: str) -> bytes:
    font_path = next((path for path in UNICODE_FONT_CANDIDATES if Path(path).exists()), None)
    if font_path is None:
        pytest.skip("Unicode font unavailable for PDF fixture")

    pdf = fitz.open()
    for text in pages:
        page = pdf.new_page(width=595, height=842)
        page.insert_textbox(
            fitz.Rect(48, 48, 547, 794),
            text,
            fontfile=font_path,
            fontname="anonimizatorunicode",
            fontsize=12,
        )
    return pdf.tobytes()


def _docx_text(data: bytes) -> str:
    document = load_document("wynik.docx", DOCX_MIME, data)
    return "\n".join(document.texts).replace("\xa0", " ")


def _digits(text: str) -> str:
    return "".join(character for character in text if character.isdigit())


FULL_DOCUMENT_PAGES = (
    (
        "Umowa dla Jan Kowalski oraz J. Kowalski. "
        "Email jan.kowalski@example.com, rachunek 41 1140 2004 0000 3102 1234 5678. "
        "Pojazd KR 7MZ18 i identyfikator host-waw-01. Dostęp z 83.21.144.9."
    ),
    (
        "Materiały obejmują korespondencję z Łódźa i dokumenty dla lokalu przy Piotrkowskiej. "
        "Dodatkowe ustalenia dotyczyły przekazania kluczy w Wrocławu przy ul. Długa 41/2. "
        "Kontakt Anna Nowak przez anna.nowak@example.com."
    ),
)

FULL_DOCUMENT_LEAKS = (
    "Jan Kowalski",
    "J. Kowalski",
    "jan.kowalski@example.com",
    "41 1140 2004 0000 3102 1234 5678",
    "KR 7MZ18",
    "host-waw-01",
    "83.21.144.9",
    "Łódźa",
    "Piotrkowskiej",
    "Wrocławu",
    "Długa",
    "Anna Nowak",
    "anna.nowak@example.com",
)


def _assert_no_known_leaks(text: str) -> None:
    leaks = [value for value in FULL_DOCUMENT_LEAKS if value in text]
    assert leaks == []
    assert "41114020040000310212345678" not in _digits(text)


def test_processor_returns_anonymized_docx_document() -> None:
    result = _process("sample.docx", DOCX_MIME, _docx_bytes("Jan Kowalski, PESEL 44051401359"))

    assert result.filename == "sample.anonimizowany.docx"
    assert result.content_type == DOCX_MIME
    assert result.data.startswith(b"PK")
    assert result.findings["PERSON"] == 1
    assert result.findings["PESEL"] == 1

    output_text = _docx_text(result.data)
    assert "Jan Kowalski" not in output_text
    assert "44051401359" not in output_text


def test_full_docx_document_regression_has_no_known_leaks() -> None:
    result = _process("pelna-umowa.docx", DOCX_MIME, _docx_bytes(*FULL_DOCUMENT_PAGES))

    assert result.filename == "pelna-umowa.anonimizowany.docx"
    assert result.content_type == DOCX_MIME
    assert result.data.startswith(b"PK")
    _assert_no_known_leaks(_docx_text(result.data))


def test_pdf_input_is_converted_to_anonymized_docx() -> None:
    data = _unicode_pdf_bytes(
        "Dane nie są fikcyjne. Zażółć gęślą jaźń. Jan Kowalski PESEL 44051401359",
        "Druga strona bez danych.",
    )

    result = _process("sample.pdf", PDF_MIME, data)

    assert result.filename == "sample.anonimizowany.docx"
    assert result.content_type == DOCX_MIME
    assert result.data.startswith(b"PK")

    output_text = _docx_text(result.data)
    assert "Zażółć gęślą jaźń" in output_text
    assert "Jan Kowalski" not in output_text
    assert "44051401359" not in output_text


def test_full_pdf_document_regression_produces_docx_without_known_leaks() -> None:
    data = _unicode_pdf_bytes(*FULL_DOCUMENT_PAGES)

    result = _process("pelna-umowa.pdf", PDF_MIME, data)

    assert result.filename == "pelna-umowa.anonimizowany.docx"
    assert result.content_type == DOCX_MIME
    assert result.data.startswith(b"PK")
    _assert_no_known_leaks(_docx_text(result.data))


def test_unsupported_text_input_is_rejected() -> None:
    with pytest.raises(DocumentError, match="DOCX i PDF"):
        _process("sample.txt", "text/plain", b"Anna Nowak email anna@example.com")


def _docx_with_comment(body: str, comment: str, author: str) -> bytes:
    parts: dict[str, bytes] = {}
    with ZipFile(BytesIO(_docx_bytes(body))) as archive:
        for name in archive.namelist():
            parts[name] = archive.read(name)
    parts["word/comments.xml"] = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:comments xmlns:w="{_W_NS}">'
        f'<w:comment w:id="1" w:author="{author}" w:initials="XX">'
        f"<w:p><w:r><w:t>{comment}</w:t></w:r></w:p>"
        f"</w:comment></w:comments>"
    ).encode()
    parts["word/people.xml"] = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w15:people xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">'
        f'<w15:person w15:author="{author}"/></w15:people>'
    ).encode()
    output = BytesIO()
    with ZipFile(output, "w") as out:
        for name, data in parts.items():
            out.writestr(name, data)
    return output.getvalue()


def test_comment_author_is_redacted_from_metadata() -> None:
    data = _docx_with_comment(
        "Zwykły tekst umowy.",
        "Uwaga od recenzenta.",
        author="Jan Kowalski",
    )

    result = _process("z-komentarzem.docx", DOCX_MIME, data)

    with ZipFile(BytesIO(result.data)) as archive:
        comments = archive.read("word/comments.xml").decode()
        people = archive.read("word/people.xml").decode()

    assert "Jan Kowalski" not in comments
    assert "Jan Kowalski" not in people
    alias = re.search(r'w:author="(Autor \d+)"', comments).group(1)
    assert f'w15:author="{alias}"' in people
    # The comment body text is still anonymized by the main pipeline.
    assert "Uwaga od recenzenta" in comments


def _docx_with_doc_props(body: str, *, author: str, last_modified_by: str) -> bytes:
    document = Document()
    document.add_paragraph(body)
    document.core_properties.author = author
    document.core_properties.last_modified_by = last_modified_by
    base = BytesIO()
    document.save(base)

    parts: dict[str, bytes] = {}
    with ZipFile(base) as archive:
        for name in archive.namelist():
            parts[name] = archive.read(name)
    parts["docProps/app.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/'
        'extended-properties">'
        f"<Manager>{author}</Manager><Company>ACME Kancelaria</Company>"
        "</Properties>"
    ).encode()
    output = BytesIO()
    with ZipFile(output, "w") as out:
        for name, data in parts.items():
            out.writestr(name, data)
    return output.getvalue()


def test_document_author_metadata_is_redacted() -> None:
    data = _docx_with_doc_props(
        "Zwykły tekst umowy.",
        author="Jan Kowalski",
        last_modified_by="Anna Nowak",
    )

    result = _process("z-metadanymi.docx", DOCX_MIME, data)

    with ZipFile(BytesIO(result.data)) as archive:
        core = archive.read("docProps/core.xml").decode()
        app = archive.read("docProps/app.xml").decode()

    assert "Jan Kowalski" not in core
    assert "Anna Nowak" not in core
    assert "<dc:creator>Autor 1</dc:creator>" in core
    assert "<cp:lastModifiedBy>Autor 2</cp:lastModifiedBy>" in core
    # The manager is the same person as the creator, so shares the alias.
    assert "Jan Kowalski" not in app
    assert "<Manager>Autor 1</Manager>" in app
    assert "ACME Kancelaria" not in app
    assert "<Company></Company>" in app


def test_processor_respects_docx_text_limit() -> None:
    with pytest.raises(DocumentError, match="przekracza limit"):
        _process("sample.docx", DOCX_MIME, _docx_bytes("abcdef"), max_text_chars=3)
