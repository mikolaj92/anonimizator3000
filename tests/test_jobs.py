import asyncio
from io import BytesIO

import pytest
from doctotext import DOCX_MIME
from docx import Document
from posejdon import ReplacementKind

from anonimizator3000.jobs import DocumentProcessingQueue, QueueRejected


def _docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class _FakeResult:
    def __init__(self, texts: list[str], findings: list[str]) -> None:
        self.texts = texts
        self.findings = findings


class _FakeAnonymizer:
    def __init__(self) -> None:
        self.seen_styles: list[ReplacementKind] = []

    def anonymize_segments(self, texts: list[str], *, replacement_style: ReplacementKind):
        self.seen_styles.append(replacement_style)
        return _FakeResult(
            [text.replace("secret", "<REDACTED>") for text in texts],
            ["TEST"],
        )


def _make_queue(anonymizer: _FakeAnonymizer, **overrides) -> DocumentProcessingQueue:
    kwargs = dict(
        anonymizer=anonymizer,
        max_text_chars=250_000,
        max_size=10,
        worker_count=1,
        max_active_jobs_per_ip=10,
        rate_limit_submissions=10,
        rate_limit_window_seconds=60,
        job_ttl_seconds=60,
    )
    kwargs.update(overrides)
    return DocumentProcessingQueue(**kwargs)


async def _wait_for_done(queue: DocumentProcessingQueue, job_id: str):
    for _ in range(100):
        snapshot = await queue.get(job_id)
        if snapshot and snapshot.status in {"done", "failed"}:
            return snapshot
        await asyncio.sleep(0.02)
    raise AssertionError("Job did not finish")


@pytest.mark.asyncio
async def test_queue_limits_active_jobs_per_ip_and_drops_source_bytes_after_processing() -> None:
    queue = _make_queue(_FakeAnonymizer(), max_active_jobs_per_ip=1)

    first = await queue.submit(
        ip="127.0.0.1",
        filename="a.docx",
        content_type=DOCX_MIME,
        data=_docx_bytes("secret text"),
    )
    with pytest.raises(QueueRejected):
        await queue.submit(
            ip="127.0.0.1",
            filename="b.docx",
            content_type=DOCX_MIME,
            data=_docx_bytes("secret text"),
        )

    await queue.start()
    try:
        done = await _wait_for_done(queue, first.id)
    finally:
        await queue.stop()

    assert done.status == "done"
    assert done.result_content_type == DOCX_MIME
    assert done.has_source_bytes is False
    assert done.findings == {"TEST": 1}

    document = await queue.result_document(first.id)
    assert document is not None
    _, content_type, data = document
    assert content_type == DOCX_MIME
    text = "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)
    assert "secret" not in text
    assert "<REDACTED>" in text


@pytest.mark.asyncio
async def test_queue_passes_replacement_style_to_anonymizer() -> None:
    anonymizer = _FakeAnonymizer()
    queue = _make_queue(anonymizer)

    job = await queue.submit(
        ip="127.0.0.1",
        filename="a.docx",
        content_type=DOCX_MIME,
        data=_docx_bytes("secret text"),
        replacement_style="labels",
    )

    await queue.start()
    try:
        done = await _wait_for_done(queue, job.id)
    finally:
        await queue.stop()

    assert done.status == "done"
    assert anonymizer.seen_styles == [ReplacementKind.CATEGORY_PLACEHOLDER]


@pytest.mark.asyncio
async def test_queue_rate_limits_submissions_per_ip() -> None:
    queue = _make_queue(_FakeAnonymizer(), rate_limit_submissions=1)

    await queue.submit(
        ip="127.0.0.1", filename="a.docx", content_type=DOCX_MIME, data=_docx_bytes("a")
    )

    with pytest.raises(QueueRejected, match="Limit uploadów"):
        await queue.submit(
            ip="127.0.0.1", filename="b.docx", content_type=DOCX_MIME, data=_docx_bytes("b")
        )
